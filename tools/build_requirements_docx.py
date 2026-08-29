from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\planner")
SOURCE = ROOT / "docs" / "自主路径探索双引擎验证需求基线.md"
OUTPUT = ROOT / "docs" / "自主路径探索双引擎验证需求基线.docx"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

BLACK = "1F2328"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "59636E"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
BORDER = "C7CDD4"


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_run_font(run, size=None, bold=None, italic=None, color=BLACK):
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_inline(paragraph, text, size=11, color=BLACK):
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, color=DARK_BLUE)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, size=9, color=MUTED)


def set_style(style, size, color, bold, before, after, line=1.1):
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line
    style.paragraph_format.keep_with_next = True


def configure_styles(doc):
    set_style(doc.styles["Normal"], 11, BLACK, False, 0, 6, 1.1)
    set_style(doc.styles["Title"], 24, BLACK, True, 0, 4, 1.0)
    set_style(doc.styles["Subtitle"], 12.5, MUTED, False, 0, 14, 1.1)
    set_style(doc.styles["Heading 1"], 16, BLUE, True, 16, 8, 1.05)
    set_style(doc.styles["Heading 2"], 13, BLUE, True, 12, 6, 1.05)
    set_style(doc.styles["Heading 3"], 12, DARK_BLUE, True, 8, 4, 1.05)
    set_style(doc.styles["List Bullet"], 11, BLACK, False, 0, 5, 1.1)
    set_style(doc.styles["List Number"], 11, BLACK, False, 0, 5, 1.1)
    for name in ("List Bullet", "List Number"):
        fmt = doc.styles[name].paragraph_format
        fmt.left_indent = Inches(0.5)
        fmt.first_line_indent = Inches(-0.25)
        fmt.tab_stops.add_tab_stop(Inches(0.5))
        fmt.keep_together = True


def configure_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]

    def create(abstract_id, num_id, fmt, text, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        p_pr.extend([tabs, ind])
        lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    next_abs = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1
    create(next_abs, next_num, "bullet", "•", "Calibri")
    bullet_num_id = next_num
    create(next_abs + 1, next_num + 1, "decimal", "%1.")
    return bullet_num_id, next_num + 1


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("自主路径探索验证  |  需求基线")
    set_run_font(run, size=9, bold=True, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5),
        WD_TAB_ALIGNMENT.RIGHT,
        WD_TAB_LEADER.SPACES,
    )
    add_inline(p, "APR-VP-REQ-001  •  V1.0", size=9, color=MUTED)
    p.add_run("\t")
    add_inline(p, "第 ", size=9, color=MUTED)
    add_page_field(p)
    add_inline(p, " 页", size=9, color=MUTED)


def add_title_block(doc):
    p = doc.add_paragraph(style="Title")
    add_inline(p, "自主路径探索双引擎验证需求基线", size=24)
    p = doc.add_paragraph(style="Subtitle")
    add_inline(p, "快速验证阶段 · 后续设计、开发、测试与验收的统一需求来源", size=12.5, color=MUTED)
    metadata = [
        ("文档编号", "APR-VP-REQ-001"),
        ("版本", "V1.0"),
        ("状态", "验证基线"),
        ("发布日期", "2026-08-14"),
        ("适用项目", "planner 自主路径探索验证项目"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{label}：")
        set_run_font(run, size=10.5, bold=True)
        run = p.add_run(value)
        set_run_font(run, size=10.5)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT)
    p_pr.append(shd)
    add_inline(
        p,
        "基线声明：本文档是验证阶段的唯一需求基线。后续实现必须引用稳定需求编号；需求变更必须同步更新版本、变更记录和验收项。",
        size=10.5,
    )


def parse_table(lines, start):
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[idx].strip().strip("|").split("|")]
        rows.append(cells)
        idx += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c) for c in rows[1]):
        rows.pop(1)
    return rows, idx


def table_widths(rows):
    cols = len(rows[0])
    if cols == 2:
        first_header = rows[0][0]
        if first_header in {"术语", "版本"}:
            return [1900, CONTENT_WIDTH_DXA - 1900]
        return [2300, CONTENT_WIDTH_DXA - 2300]
    if cols == 3:
        return [1500, 1800, CONTENT_WIDTH_DXA - 3300]
    base = CONTENT_WIDTH_DXA // cols
    widths = [base] * cols
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = table_widths(rows)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_borders(cell)
            if r_idx == 0:
                shade_cell(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_inline(p, text, size=9.5 if len(text) > 80 else 10, color=BLACK)
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)
    after.paragraph_format.space_before = Pt(2)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    bullet_num_id, decimal_num_id = configure_numbering(doc)
    add_header_footer(doc)
    add_title_block(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = 0
    while start < len(lines) and not lines[start].startswith("## 1."):
        start += 1
    idx = start
    while idx < len(lines):
        raw = lines[idx].rstrip()
        stripped = raw.strip()
        if not stripped or stripped == "---":
            idx += 1
            continue
        if stripped.startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if heading:
            level = min(len(heading.group(1)) - 1, 3)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, heading.group(2), size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK_BLUE)
            idx += 1
            continue
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.right_indent = Inches(0.18)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), CALLOUT)
            p_pr.append(shd)
            add_inline(p, stripped.lstrip("> "), size=10.5)
            idx += 1
            continue
        bullet = re.match(r"^-\s+(.*)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            apply_numbering(p, bullet_num_id)
            add_inline(p, bullet.group(1), size=11)
            idx += 1
            continue
        if numbered:
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, decimal_num_id)
            add_inline(p, numbered.group(1), size=11)
            idx += 1
            continue
        p = doc.add_paragraph()
        p.paragraph_format.keep_together = False
        add_inline(p, stripped, size=11)
        idx += 1

    props = doc.core_properties
    props.title = "自主路径探索双引擎验证需求基线"
    props.subject = "Anytime A* 与遗传算法并行独立对照的快速验证需求"
    props.author = "planner project"
    props.keywords = "自主路径探索, Anytime A*, 遗传算法, PathSimulator, 需求基线"
    props.comments = "Generated from the version-controlled Markdown requirements baseline."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
